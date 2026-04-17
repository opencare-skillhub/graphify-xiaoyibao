from __future__ import annotations

import fnmatch
import json
import os
import re
from enum import Enum
from pathlib import Path


class FileType(str, Enum):
    CODE = "code"
    DOCUMENT = "document"
    PAPER = "paper"
    IMAGE = "image"
    VIDEO = "video"


_MANIFEST_PATH = "graphify-out/manifest.json"

CODE_EXTENSIONS = {
    ".py", ".ts", ".js", ".jsx", ".tsx", ".go", ".rs", ".java", ".cpp", ".cc", ".cxx", ".c",
    ".h", ".hpp", ".rb", ".swift", ".kt", ".kts", ".cs", ".scala", ".php", ".lua", ".toc",
    ".zig", ".ps1", ".ex", ".exs", ".m", ".mm", ".jl", ".vue", ".svelte", ".dart", ".v", ".sv",
}
DOC_EXTENSIONS = {".md", ".txt", ".rst"}
PAPER_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".heic", ".heif"}
OFFICE_EXTENSIONS = {".docx", ".xlsx"}
DICOM_EXTENSIONS = {".dcm", ".dicom"}
VIDEO_EXTENSIONS = {".mp4", ".mov", ".webm", ".mkv", ".avi", ".m4v", ".mp3", ".wav", ".m4a", ".ogg"}

CORPUS_WARN_THRESHOLD = 50_000
CORPUS_UPPER_THRESHOLD = 500_000
FILE_COUNT_UPPER = 200

_SENSITIVE_PATTERNS = [
    re.compile(r"(^|[\\/])\.(env|envrc)(\.|$)", re.IGNORECASE),
    re.compile(r"\.(pem|key|p12|pfx|cert|crt|der|p8)$", re.IGNORECASE),
    re.compile(r"(credential|secret|passwd|password|token|private_key)", re.IGNORECASE),
    re.compile(r"(id_rsa|id_dsa|id_ecdsa|id_ed25519)(\.pub)?$"),
    re.compile(r"(\.netrc|\.pgpass|\.htpasswd)$", re.IGNORECASE),
    re.compile(r"(aws_credentials|gcloud_credentials|service.account)", re.IGNORECASE),
]

_PAPER_SIGNALS = [
    re.compile(r"\barxiv\b", re.IGNORECASE),
    re.compile(r"\bdoi\s*:", re.IGNORECASE),
    re.compile(r"\babstract\b", re.IGNORECASE),
    re.compile(r"\bproceedings\b", re.IGNORECASE),
    re.compile(r"\bjournal\b", re.IGNORECASE),
    re.compile(r"\bpreprint\b", re.IGNORECASE),
    re.compile(r"\\cite\{"),
    re.compile(r"\[\d+\]"),
    re.compile(r"\[\n\d+\n\]"),
    re.compile(r"eq\.\s*\d+|equation\s+\d+", re.IGNORECASE),
    re.compile(r"\d{4}\.\d{4,5}"),
    re.compile(r"\bwe propose\b", re.IGNORECASE),
    re.compile(r"\bliterature\b", re.IGNORECASE),
]
_PAPER_SIGNAL_THRESHOLD = 3

_ASSET_DIR_MARKERS = {".imageset", ".xcassets", ".appiconset", ".colorset", ".launchimage"}
_SKIP_DIRS = {
    "venv", ".venv", "env", ".env", "node_modules", "__pycache__", ".git",
    "dist", "build", "target", "out", "site-packages", "lib64", "graphify-out",
    ".pytest_cache", ".mypy_cache", ".ruff_cache", ".tox", ".eggs", "*.egg-info",
}
_SKIP_FILES = {
    "package-lock.json", "yarn.lock", "pnpm-lock.yaml", "Cargo.lock", "poetry.lock",
    "Gemfile.lock", "composer.lock", "go.sum", "go.work.sum",
}

_MEDICAL_DIRECTORY_BUCKETS = {
    "01_基础信息": "basic_info",
    "02_确诊信息": "diagnosis",
    "03_基因与病理详情": "genetics_pathology",
    "04_治疗记录": "treatment",
    "05_影像资料": "imaging",
    "06_检验指标与曲线": "labs_markers",
    "07_用药方案与提醒": "medication",
    "08_并发症预防与风险管理": "risk_management",
    "09_营养评估": "nutrition",
    "10_心理评估": "psychology",
    "11_随访与复发监测": "follow_up",
    "12_其他": "other",
}


def _medical_bucket(path: Path) -> str | None:
    for part in path.parts:
        if part in _MEDICAL_DIRECTORY_BUCKETS:
            return _MEDICAL_DIRECTORY_BUCKETS[part]
    return None


def medical_bucket_for_path(path: str | Path) -> str | None:
    p = Path(path)
    return _medical_bucket(p)


def summarize_medical_layout(paths: list[str | Path]) -> dict[str, int]:
    hits = {value: 0 for value in _MEDICAL_DIRECTORY_BUCKETS.values()}
    for raw in paths:
        bucket = medical_bucket_for_path(raw)
        if bucket:
            hits[bucket] += 1
    return {k: v for k, v in hits.items() if v > 0}



def _is_sensitive(path: Path) -> bool:
    name = path.name
    full = str(path)
    return any(p.search(name) or p.search(full) for p in _SENSITIVE_PATTERNS)


def _looks_like_paper(path: Path) -> bool:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")[:3000]
        hits = sum(1 for pattern in _PAPER_SIGNALS if pattern.search(text))
        return hits >= _PAPER_SIGNAL_THRESHOLD
    except Exception:
        return False


def classify_file(path: Path) -> FileType | None:
    if path.name.lower().endswith(".blade.php"):
        return FileType.CODE
    ext = path.suffix.lower()
    if ext in CODE_EXTENSIONS:
        return FileType.CODE
    if ext in PAPER_EXTENSIONS:
        if any(part.endswith(tuple(_ASSET_DIR_MARKERS)) for part in path.parts):
            return None
        return FileType.PAPER
    if ext in IMAGE_EXTENSIONS:
        return FileType.IMAGE
    if ext in DICOM_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in DOC_EXTENSIONS:
        if _looks_like_paper(path):
            return FileType.PAPER
        return FileType.DOCUMENT
    if ext in OFFICE_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in VIDEO_EXTENSIONS:
        return FileType.VIDEO
    return None


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception:
        return ""


def docx_to_markdown(path: Path) -> str:
    try:
        from docx import Document

        doc = Document(str(path))
        lines: list[str] = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except Exception:
        return ""


def xlsx_to_markdown(path: Path) -> str:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[list[str]] = []
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            sections.extend([header, sep])
            for row in rows[1:]:
                sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except Exception:
        return ""


def convert_office_file(path: Path, out_dir: Path) -> Path | None:
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None
    if not text.strip():
        return None
    import hashlib

    out_dir.mkdir(parents=True, exist_ok=True)
    name_hash = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    out_path.write_text(f"<!-- converted from {path.name} -->\n\n{text}", encoding="utf-8")
    return out_path


def count_words(path: Path) -> int:
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return len(extract_pdf_text(path).split())
        if ext == ".docx":
            return len(docx_to_markdown(path).split())
        if ext == ".xlsx":
            return len(xlsx_to_markdown(path).split())
        return len(path.read_text(encoding="utf-8", errors="ignore").split())
    except Exception:
        return 0


def _is_noise_dir(part: str) -> bool:
    if part in _SKIP_DIRS:
        return True
    if part == "archive":
        return True
    if part.endswith("_venv") or part.endswith("_env"):
        return True
    if part.endswith(".egg-info"):
        return True
    return False


def _load_graphifyignore(root: Path) -> list[tuple[Path, str]]:
    patterns: list[tuple[Path, str]] = []
    current = root.resolve()
    while True:
        ignore_file = current / ".graphifyignore"
        if ignore_file.exists():
            for line in ignore_file.read_text(encoding="utf-8", errors="ignore").splitlines():
                line = line.strip()
                if line and not line.startswith("#"):
                    patterns.append((current, line))
        if (current / ".git").exists():
            break
        parent = current.parent
        if parent == current:
            break
        current = parent
    return patterns


def _is_ignored(path: Path, root: Path, patterns: list[tuple[Path, str]]) -> bool:
    if not patterns:
        return False

    def _matches(rel: str, pattern: str) -> bool:
        parts = rel.split("/")
        if fnmatch.fnmatch(rel, pattern):
            return True
        if fnmatch.fnmatch(path.name, pattern):
            return True
        for i, part in enumerate(parts):
            if fnmatch.fnmatch(part, pattern):
                return True
            if fnmatch.fnmatch("/".join(parts[: i + 1]), pattern):
                return True
        return False

    for anchor, pattern in patterns:
        p = pattern.strip("/")
        if not p:
            continue
        try:
            rel = str(path.relative_to(root)).replace(os.sep, "/")
            if _matches(rel, p):
                return True
        except ValueError:
            pass
        if anchor != root:
            try:
                rel_anchor = str(path.relative_to(anchor)).replace(os.sep, "/")
                if _matches(rel_anchor, p):
                    return True
            except ValueError:
                pass
    return False


def detect(root: Path, *, follow_symlinks: bool = False) -> dict:
    files: dict[FileType, list[str]] = {
        FileType.CODE: [],
        FileType.DOCUMENT: [],
        FileType.PAPER: [],
        FileType.IMAGE: [],
        FileType.VIDEO: [],
    }
    total_words = 0
    skipped_sensitive: list[str] = []
    medical_directory_hits = {value: 0 for value in _MEDICAL_DIRECTORY_BUCKETS.values()}
    ignore_patterns = _load_graphifyignore(root)

    memory_dir = root / "graphify-out" / "memory"
    scan_paths = [root]
    if memory_dir.exists():
        scan_paths.append(memory_dir)

    seen: set[Path] = set()
    all_files: list[Path] = []
    for scan_root in scan_paths:
        in_memory_tree = memory_dir.exists() and str(scan_root).startswith(str(memory_dir))
        for dirpath, dirnames, filenames in os.walk(scan_root, followlinks=follow_symlinks):
            dp = Path(dirpath)
            if follow_symlinks and os.path.islink(dirpath):
                real = os.path.realpath(dirpath)
                parent_real = os.path.realpath(os.path.dirname(dirpath))
                if parent_real == real or parent_real.startswith(real + os.sep):
                    dirnames.clear()
                    continue
            if not in_memory_tree:
                dirnames[:] = [
                    d for d in dirnames
                    if not d.startswith(".")
                    and not _is_noise_dir(d)
                    and not _is_ignored(dp / d, root, ignore_patterns)
                ]
            for fname in filenames:
                if fname in _SKIP_FILES:
                    continue
                p = dp / fname
                if p not in seen:
                    seen.add(p)
                    all_files.append(p)

    converted_dir = root / "graphify-out" / "converted"

    for p in all_files:
        in_memory = memory_dir.exists() and str(p).startswith(str(memory_dir))
        if not in_memory:
            if p.name.startswith("."):
                continue
            if str(p).startswith(str(converted_dir)):
                continue
        if _is_ignored(p, root, ignore_patterns):
            continue
        if _is_sensitive(p):
            skipped_sensitive.append(str(p))
            continue
        ftype = classify_file(p)
        if ftype:
            bucket = _medical_bucket(p.relative_to(root) if p.is_absolute() else p)
            if bucket:
                medical_directory_hits[bucket] += 1
            if p.suffix.lower() in OFFICE_EXTENSIONS:
                md_path = convert_office_file(p, converted_dir)
                if md_path:
                    files[ftype].append(str(md_path))
                    total_words += count_words(md_path)
                else:
                    skipped_sensitive.append(str(p) + " [office conversion failed - install office deps]")
                continue
            files[ftype].append(str(p))
            if ftype != FileType.VIDEO:
                total_words += count_words(p)

    total_files = sum(len(v) for v in files.values())
    needs_graph = total_words >= CORPUS_WARN_THRESHOLD
    warning: str | None = None
    if not needs_graph:
        warning = f"Corpus is ~{total_words:,} words - fits in a single context window. You may not need a graph."
    elif total_words >= CORPUS_UPPER_THRESHOLD or total_files >= FILE_COUNT_UPPER:
        warning = (
            f"Large corpus: {total_files} files · ~{total_words:,} words. "
            f"Semantic extraction will be expensive. Consider running on a subfolder."
        )

    return {
        "files": {k.value: v for k, v in files.items()},
        "total_files": total_files,
        "total_words": total_words,
        "needs_graph": needs_graph,
        "warning": warning,
        "skipped_sensitive": skipped_sensitive,
        "graphifyignore_patterns": len(ignore_patterns),
        "medical_directory_hits": {k: v for k, v in medical_directory_hits.items() if v > 0},
    }


def load_manifest(manifest_path: str = _MANIFEST_PATH) -> dict[str, float]:
    try:
        return json.loads(Path(manifest_path).read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_manifest(files: dict[str, list[str]], manifest_path: str = _MANIFEST_PATH) -> None:
    manifest: dict[str, float] = {}
    for file_list in files.values():
        for f in file_list:
            try:
                manifest[f] = Path(f).stat().st_mtime
            except OSError:
                pass
    Path(manifest_path).parent.mkdir(parents=True, exist_ok=True)
    Path(manifest_path).write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def detect_incremental(root: Path, manifest_path: str = _MANIFEST_PATH) -> dict:
    full = detect(root)
    manifest = load_manifest(manifest_path)
    if not manifest:
        full["incremental"] = True
        full["new_files"] = full["files"]
        full["unchanged_files"] = {k: [] for k in full["files"]}
        full["new_total"] = full["total_files"]
        return full

    new_files: dict[str, list[str]] = {k: [] for k in full["files"]}
    unchanged_files: dict[str, list[str]] = {k: [] for k in full["files"]}
    for ftype, file_list in full["files"].items():
        for f in file_list:
            stored_mtime = manifest.get(f)
            try:
                current_mtime = Path(f).stat().st_mtime
            except Exception:
                current_mtime = 0
            if stored_mtime is None or current_mtime > stored_mtime:
                new_files[ftype].append(f)
            else:
                unchanged_files[ftype].append(f)

    current_files = {f for flist in full["files"].values() for f in flist}
    deleted_files = [f for f in manifest if f not in current_files]

    new_total = sum(len(v) for v in new_files.values())
    full["incremental"] = True
    full["new_files"] = new_files
    full["unchanged_files"] = unchanged_files
    full["new_total"] = new_total
    full["deleted_files"] = deleted_files
    return full
